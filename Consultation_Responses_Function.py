


def responses_in_word(data, directory, xlsxfile, GN_name, name_column, altname_column, firstcol, lastcol, rankcolumns, matrixcolumns,selectall_inmatrix):
    import os
    import pandas as pd
    import matplotlib.pyplot as plt
    from matplotlib import cm
    import numpy as np
    import re
    
    
    os.chdir(directory)
    #data =  pd.read_excel(xlsxfile)
    
    
    graphcols= []
    for i in range(data.shape[1]):
        data2 = data[[data.columns[i]]].assign(Responses=[1]*data.shape[0]).groupby(data.columns[i]).sum()
        if data2.shape[0]<=5:
            graphcols = graphcols + [i]
                
    
    
    def smartplotfunction(dat,i, data):#dat is a single column dataset. Like this: data[[data.columns[i]]]
            dat2 = dat.assign(Responses=[1]*data.shape[0]).groupby(data.columns[i]).sum()
            if 'Agree' in dat2.Responses.index:
                dat2 = dat2.reindex(["Strongly agree", "Agree","Neutral","Disagree", "Strongly disagree"])
                cs= cm.RdBu(np.arange(dat2.shape[0])/dat2.shape[0])
            elif 'Undecided' in dat2.Responses.index:
                dat2 = dat2.reindex(["Yes", "No","Undecided"])
                cs=cm.Set3(np.arange(dat2.shape[0])/dat2.shape[0])
            elif not('Undecided' in dat2.Responses.index) and 'Yes' in dat2.Responses.index and 'No' in dat2.Responses.index:
                dat2 = dat2.reindex(["Yes", "No"])
                cs=cm.Set3(np.arange(dat2.shape[0])/dat2.shape[0])
            else:
                cs=cm.Set3(np.arange(dat2.shape[0])/dat2.shape[0])
            
            dat2 = dat2.fillna(0)
            dat2.Responses = dat2.Responses.astype(int)
            total = sum(dat2.Responses)
            plt.pie(dat2.Responses, labels = dat2.index, colors=cs, autopct=lambda p: '{:.0f}'.format(p * total / 100))
            #plt.legend(dat2.index, loc='lower right')
            plt.savefig("fig" + str(i) +".png")
            plt.close()
    
    def plotfunction(dat,i,data):#this wants a single column dataset. Like this: data[[data.columns[i]]]
            dat2 = dat.assign(Responses=[1]*data.shape[0]).groupby(data.columns[i]).sum()
            plt.pie(dat2.Responses, labels = dat2.index)
            #plt.legend(dat2.index, loc='lower right')
            plt.savefig("fig" + str(i) +".png")
            plt.close()    
    
    
    #data_cleanup to simply remove full non-responses and responses that did not provide identification.
    keep = [True]*data.shape[0]
    for i in range(data.shape[0]):
        if (data.iloc[i,firstcol:lastcol].isnull()==False).sum()<=2:
            keep[i]= False
        if str(data.iloc[i, name_column])=='nan' and str(data.iloc[i, name_column])=='nan': #delete if they have neither of the name columns
            keep[i] = False
        if str(data.iloc[i, name_column])=='nan' and str(data.iloc[i, name_column])!='nan':
            data.iloc[i, name_column] = data.iloc[i, altname_column]
            
    data = data[keep]
    
    
    
    #rank type questions:
    for i in range(len(rankcolumns)):
        rankcols = rankcolumns[i] #ith set of rank columns
        
        x = []
        y = []
        
        pattern = re.compile("([0-9]|[1-9][0-9])\.\s")
        k=0
        for j in rankcols:
            topic = data.columns[j].split("[Topic:")[1].split('; Category')[0][0:]
            x = x + [topic]
            y.append((data.iloc[:,j]==1).sum())
            for l in range(data.shape[0]):
                if str(data.iloc[l,j])!= "nan":
                    data.iloc[l,j] = str(int(data.iloc[l,j]))+ ": " + topic
            if k>0:
                data.iloc[:,j-k] = data.iloc[:,j-k] + " " + data.iloc[:,j]
            else:
                cols = list(data.columns)
                question = data.columns[j].split("Question:")[1]
                question = re.split(pattern, data.columns[j].split("[Topic:")[0])[0] + question[:-1]
                cols[j] = question
                data.columns = cols
            k = k+1
    
        plt.bar(x, y, color='royalblue', alpha=0.7)
        plt.grid(color='#95a5a6', linestyle='--', linewidth=2, axis='y', alpha=0.7)
        plt.ylabel('Number of times ranked 1st')
        plt.xticks(rotation=90)
        plt.tight_layout()
        plt.savefig("fig" + str(rankcols[0]) +".png")
        plt.close()
    
    
    #make one list out of the list of list. remove the first item of each.
    #we will skip the columns in this list later.
    other_rankcolumns = []  
    rankcolumns2 = []        
    for i in range(len(rankcolumns)):
        other_rankcolumns = other_rankcolumns + rankcolumns[i][1:]
        rankcolumns2 = rankcolumns2 + rankcolumns[i]
    
    rankcolumns = rankcolumns2
    
    #plots for select all that apply cols
    #identify first all sets of columns that refer to a select all that apply question.
    selectallthatapplycols_listoflists = []
    sublistcounter = 0
    sublist = []
    question0 = ""
    for i in range(firstcol, lastcol):
        if '[Question:' in data.columns[i] and not('Specify [Question:' in data.columns[i]):
            pattern = re.compile("([0-9]|[1-9][0-9])\.\s")
            question = data.columns[i].split("[Question:")[1]
            question = re.split(pattern, data.columns[i].split("[Question:")[0])[1] +"." + question[:-1]
            if question0 != question:#this is intended to be the first item
                if sublistcounter>1:#save the old list in the list of list
                    selectallthatapplycols_listoflists = selectallthatapplycols_listoflists  + [sublist]
                sublist = [i]
                sublistcounter = sublistcounter+1
                question0 = question
            else:
                sublist = sublist + [i]
            
            
    #save the last sublist
    selectallthatapplycols_listoflists = selectallthatapplycols_listoflists  + [sublist]
    #--------------------------------------------------------
    

    selectallthatapplyoptions = [""]*data.shape[1]
    # go through this list of lists and create graphs:
    for listi in selectallthatapplycols_listoflists:
        if len(listi)>0:
            colsi = data.columns[listi]
            data2 = data[colsi]
            question = data.columns[listi[0]].split("[Question:")[1]
            question = re.split(pattern, data.columns[listi[0]].split("[Question:")[0])[1] +"." + question[:-1]
            x = []
            y = []
            for i in range(data2.shape[1]):
                data3 = data2.iloc[:,i].dropna()
                x.append(list(data3[data3!='-'])[0])
                y.append(len(list(data3[data3!='-'])))
                
                selectallthatapplyoption = data2.columns[i].split("[Question:")[0]
                pattern = re.compile("([0-9]|[1-9][0-9])\.\s")
                selectallthatapplyoption = re.split(pattern, selectallthatapplyoption)[2]
                selectallthatapplyoptions[listi[0]] = selectallthatapplyoptions[listi[0]] + str(i+1)+'. ' + selectallthatapplyoption
                if i >0:
                    data.iloc[:,listi[0]] = data.iloc[:,listi[0]] + "; " + data2.iloc[:,i] 
            
            plt.bar(x, y, color='royalblue', alpha=0.7)
            plt.grid(color='#95a5a6', linestyle='--', linewidth=2, axis='y', alpha=0.7)
            plt.ylabel('Number of times selected')
            plt.xticks(rotation=90)
            plt.savefig("fig" + str(listi[0]) +".png", bbox_inches='tight')
            plt.close()
        
            renamecolumns = list(data.columns)
            renamecolumns[listi[0]] = question
            data.columns = renamecolumns

    
        
    
    additional_selectallthatapply_columns = []
    for listi in selectallthatapplycols_listoflists:
        if len(listi)>0:
            additional_selectallthatapply_columns = additional_selectallthatapply_columns + listi[1:]
    
    
    #Deal with please specify questions - this assumes there is only one Specify per question
    pleasespecify_columns = []
    for i in range(firstcol, lastcol):
        if 'Specify [Question:' in data.columns[i]:
            pleasespecify_columns = pleasespecify_columns + [i]
    
        if 'Specify [Question:' in data.columns[i +1]:
            
            #graph
            smartplotfunction(data[[data.columns[i]]],i, data)
            
            for j in range(data.shape[0]):
                if str(data.iloc[j, i+1])!="nan":
                    data.iloc[j,i] = data.iloc[j,i] + ": " + data.iloc[j,i+1]
                    
    
    #this is for matrix questions that have select all that apply questions
    othermatrixcolumns = []
    for i in range(len(matrixcolumns)):
        if selectall_inmatrix[i]!="":
            for j in matrixcolumns[i]:
                category = data.columns[j].split("; Category:")[1].split("; Question:")[0]
                if selectall_inmatrix[i]==category and not(j in othermatrixcolumns): #this only deals with the select all that apply columns 
                    x = []
                    y = []
                    data3 = data.iloc[:,j].dropna()
                    x.append(list(data3[data3!='-'])[0])
                    y.append(len(list(data3[data3!='-'])))
                
                    pattern = re.compile("([0-9]|[1-9][0-9])\.\s")  
                    question = re.split(pattern, data.columns[j])[1] + ". " + data.columns[j].split("; Question:")[1][:-1]
                    question = question + " " + data.columns[j].split("[Topic:")[1].split("; ")[0]
                    columnnames = list(data.columns)
                    columnnames[j] = question
                    data.columns = columnnames
                    
                    #after identifying the question (that includes the topic) we go through the rest of the matrix and identify the columns
                    #that have the same question and are still in the select all that apply category
                    for k in range(j+1, matrixcolumns[i][len(matrixcolumns[i])-1]+1):
                        pattern = re.compile("([0-9]|[1-9][0-9])\.\s")  
                        question = re.split(pattern, data.columns[k])[1] + ". " + data.columns[k].split("; Question:")[1][:-1]
                        question = question + " " + data.columns[k].split("[Topic:")[1].split("; ")[0]
                        category = data.columns[k].split("; Category:")[1].split("; Question:")[0]
                        if question == data.columns[j] and selectall_inmatrix[i]==category:
                            data3 = data.iloc[:,k].dropna()
                            if data3[data3!='-'].shape[0]>0:
                                x.append(list(data3[data3!='-'])[0])
                                y.append(len(list(data3[data3!='-'])))
                            
                            othermatrixcolumns= othermatrixcolumns + [k]
                            for l in range(data.shape[0]):
                                if str(data.iloc[l,k])!="nan" and str(data.iloc[l,k])!="-":
                                    data.iloc[l,j] = str(data.iloc[l,j]) + "; " + str(data.iloc[l,k])
                        if question == data.columns[j] and selectall_inmatrix[i]!=category:
                            data3 = data.iloc[:,k].dropna()
                            x.append(category)
                            y.append(len(list(data3[data3!='-'])))
                            othermatrixcolumns= othermatrixcolumns + [k]
                            for l in range(data.shape[0]):
                                if str(data.iloc[l,k])!="nan":
                                    data.iloc[l,j] = str(data.iloc[l,j]) + "\n" + category + ": " + str(data.iloc[l,k])
                    
                    plt.bar(x, y, color='royalblue', alpha=0.7)
                    plt.grid(color='#95a5a6', linestyle='--', linewidth=2, axis='y', alpha=0.7)
                    plt.ylabel('Number of times selected')
                    plt.xticks(rotation=90)
                    plt.savefig("fig" + str(j) +".png", bbox_inches='tight')
                    plt.close()
    
    
    
    
    #other matrix questions:
    n = 0
    for listi in matrixcolumns:
        if selectall_inmatrix[n]=="":
            if len(listi)>0: #only if there is actually at least one matrix question
                colsi = data.columns[listi]
                data2 = data[colsi]
                question = data2.columns[0].split("; Question:")[1]
                pattern = re.compile("([0-9]|[1-9][0-9])\.\s")
                if not(re.search(pattern, data2.columns[0].split("[Question:")[0]) is None):
                    question = re.split(pattern, data2.columns[0].split("[Question:")[0])[1] +". " + question[:-1]
                else:
                    pattern = re.compile("([0-9]|[1-9][0-9])B.\s")
                    question = re.split(pattern, data2.columns[0].split("[Question:")[0])[1] +". " + question[:-1]
                topiclist = [""]*len(listi)
                categorylist = [""]*len(listi)
                for i in range(len(listi)):
                    topiclist[i] = data2.columns[i].split("[Topic:")[1].split("; ")[0]
                    categorylist[i] = data2.columns[i].split("; Category:")[1].split("; Question:")[0]
                    if len(categorylist[i])==1:
                        categorylist[i] = ""
                
                i =0
                while not(listi[i] in othermatrixcolumns):
                    topic = topiclist[i]
                    #rename column
                    columntext = question + " " + topic
                    columnnames = list(data.columns)
                    columnnames[listi[i]] = columntext
                    data.columns = columnnames
                    
                    k=0
                    l ="" 
                    for j in range(len(listi)):
                        #loop through the matrix, to get the columns that belong together
                        if topiclist[j] == topic:
                            if k >0:
                                othermatrixcolumns = othermatrixcolumns + [listi[j]]
                            k = k+1
                            if l=="":
                                l = j #this only changes to j the first time the topic is correct
                            if listi[j] in graphcols:
                                smartplotfunction(data[[data.columns[listi[j]]]],listi[j], data)
                            
                            for m in range(data.shape[0]):
                                if j ==l:
                                    data.iloc[m,listi[l]] = categorylist[j] + str(data.iloc[m,listi[j]])
                                else:
                                    if str(data.iloc[m,listi[j]])!="nan":
                                        data.iloc[m,listi[l]] = data.iloc[m,listi[l]] + "\n" + categorylist[j] + " " + data.iloc[m,listi[j]]
                    
                    i = i+1
        n = n+1
    
    matrixcolumns2 = []        
    for i in range(len(matrixcolumns)):
        matrixcolumns2 = matrixcolumns2 + matrixcolumns[i]
        
    
    matrixcolumns = matrixcolumns2
    
    
    B_columns = []
    def smartcolumnsfunction(data):
        B_columns = []
        pattern = re.compile("([0-9]|[1-9][0-9])B.\s") #this looks for any number from 0-99 followed by "B. "
        for i in range(data.shape[1]):
            if pattern.match(data.columns[i]) and not(i in matrixcolumns):
                B_columns = B_columns + [i]
        return B_columns
    
    B_columns  = smartcolumnsfunction(data)
    
    
    #this function is 0 when there is no "B" column that belongs to there i'th columns.
    #if there is a "B" column that belongs to the i'th column it returns how many columns away it is
    def next_B_column(i): 
        B_column_k_away = 0
        j = 1   
        if (i+1) in B_columns:
            B_column_k_away = 1 #this means the B_column is in the next column
        if (i+1) in  other_rankcolumns:
            while (i+j) in other_rankcolumns:
                B_column_k_away = j
                j =j+1
            if (i+B_column_k_away+1) in B_columns:
                B_column_k_away = B_column_k_away+1 #this means the B_column is in the column after a set of rank question columns
            else:
                B_column_k_away = 0
        if (i+1) in pleasespecify_columns:
            while (i+j) in pleasespecify_columns:
                B_column_k_away = j
                j =j+1
            if (i+B_column_k_away+1) in B_columns: 
                B_column_k_away = B_column_k_away+1 #this means the B_column is in the column after a set of please specify columns
            else:
                B_column_k_away = 0
        if (i+1) in additional_selectallthatapply_columns:
            while (i+j) in additional_selectallthatapply_columns:
                B_column_k_away = j
                j =j+1
            if (i+B_column_k_away+1) in B_columns: 
                B_column_k_away = B_column_k_away+1 #this means the B_column is in the column after a set of select all that apply columns
            else:
                B_column_k_away = 0
            
        return B_column_k_away
    

    
    
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    document = Document()
    
    
    paragraph = document.add_heading("\nResponses to the Global Consultation of:\n "+ GN_name, 1)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    para = document.add_paragraph()
    para.add_run("A total of " + str(data.shape[0]) + " respondents contributed to this global consultation. Completely anonymous contributions are excluded.")
    
    for i in range(firstcol,(lastcol+1)):
        if i in B_columns or i in additional_selectallthatapply_columns or i in pleasespecify_columns or i in other_rankcolumns or i in othermatrixcolumns: 
            continue
    
    
        nextB = next_B_column(i)
    
    
        if nextB>0:#if the next question is a "B." question
            document.add_heading(data.columns[i] +" " + data.columns[i+nextB], 3)
        else:
            document.add_heading(data.columns[i], 3)
        
        
        
        if selectallthatapplyoptions[i] != "":
            document.add_heading(selectallthatapplyoptions[i], 4)
        
        if i in graphcols and selectallthatapplyoptions[i] == "" and not((i+1) in pleasespecify_columns) and not(i in rankcolumns) and not(i in matrixcolumns): #this excludes select all that apply questions
            smartplotfunction(data[[data.columns[i]]],i, data)
            
        
        if i in graphcols and not(i in othermatrixcolumns): #this excludes select all that apply questions
    
            document.add_picture("fig" + str(i) +".png")
        
        
        if nextB>0:
            for j in range(data.shape[0]):
                if pd.isna(data.iloc[j,i])==False or pd.isna(data.iloc[j,i+nextB])==False:
                    para = document.add_paragraph()
                    bold_para = para.add_run(str(data.iloc[j, name_column])+ ": ")
                    bold_para.bold = True
                    if pd.isna(data.iloc[j,i])==False:
                        para.add_run(str(data.iloc[j,i]))
                    if pd.isna(data.iloc[j,i+nextB])==False:
                        para.add_run("\n"+ str(data.iloc[j,i+nextB]))       
            i = i+nextB            
        else:
            for j in range(data.shape[0]):
                if pd.isna(data.iloc[j,i])==False:
                    para = document.add_paragraph()
                    bold_para = para.add_run(str(data.iloc[j, name_column])+ ": ")
                    bold_para.bold = True
                    para.add_run(str(data.iloc[j,i]))
                    
    document.save('Global Consultation - '+ GN_name +' .docx')

